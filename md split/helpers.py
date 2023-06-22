
"""
A set of helper functions to convert a word document
of Welikia Gazetteer entries into individual Markdown files

AUTHOR: Lucinda Royte
DATE: 3.6.23

"""

import docx
import re

def find_start_text_index(doc_file, start_index=0):
    """
    Returns the index of the first occurrence of "<start_text>"
    in a given word document starting from the given index

    Parameters: word document, int
    Preconditions: 
    """

    # Open the Word document
    doc = docx.Document(doc_file)

    # Iterate over all paragraphs starting from the given index
    for i in range(start_index+1, len(doc.paragraphs)):
        # Check if <start_text> appears in the paragraph text
        if '<start_text>' in doc.paragraphs[i].text:
            # Return the index of the paragraph containing <start_text>
            return i



def get_placename(doc_file,start_index):
    """
    Returns string two lines above given phrase "<start_text>"

    Parameters: word document, int
    Preconditions:
    """

    # Open the Word document
    doc = docx.Document(doc_file)

    entry_title = doc.paragraphs[start_index-2].text
    placename = trim_placename(entry_title)
    return placename



def trim_placename(text):
    """ 
    Returns placename stripped of parenthese and special characters

    Parameters: string
    Preconditions: 
    """
    # Remove text within parentheses
    start_index = text.find("(")
    if start_index != -1:
        end_index = text.find(")", start_index)
        if end_index != -1:
            text = text[:start_index].strip() + text[end_index+1:].strip()

    # Remove special characters
    trimmed_placename = re.sub('[^A-Za-z0-9\s]+', '', text)

    return trimmed_placename.strip()



def find_end_text_index(doc_file, start_index):
    # Open the Word document
    doc = docx.Document(doc_file)

    # Iterate over all paragraphs starting from the given index
    for i in range(start_index, len(doc.paragraphs)):
        # Check if <end_text> appears in the paragraph text
        if '<end_text>' in doc.paragraphs[i].text:
            # Return the index of the paragraph containing <end_text>
            return i

    # If no <end_text> is found, return None
    return None



def get_placename_id(doc_file, end_index):
    # Open the Word document
    doc = docx.Document(doc_file)

    # Get the text of the paragraph containing <end_text>
    end_text = doc.paragraphs[end_index].text
    # Strip line from placename id int
    placename_id = end_text.strip("<end_text> id=")
    return placename_id


def get_entry_text(doc_file, start_index, end_index): #docx to md after edits made and accepted in big docx
    # Open the Word document
    doc = docx.Document(doc_file)

    # Initialize an empty string to hold the entry text
    entry_text = ""

    # Iterate over the paragraphs between the start and end indexes
    for i in range(start_index + 1, end_index):
        # Get the paragraph object
        para = doc.paragraphs[i]

        # Initialize an empty list to hold the formatted text
        formatted_runs = []

        # Iterate over the runs in the paragraph
        for run in para.runs:
            # Check if the run is bold
            if run.bold:
                formatted_runs.append('**' + run.text + '**')
            # Check if the run is italic
            elif run.italic:
                formatted_runs.append('*' + run.text + '*')
            else:
                formatted_runs.append(run.text)

        # Join the formatted runs and add a newline character
        formatted_text = ''.join(formatted_runs) + '\n'

        # Append the formatted text to the entry text
        entry_text += formatted_text

    # Return the entry text
    return entry_text.strip()


# import docx
# from docx.enum.text import WD_COLOR_INDEX

# def get_entry_text(doc_file, start_index, end_index): #docx to docx
#     # Open the Word document
#     doc = docx.Document(doc_file)

#     # Get the entry paragraphs
#     entry_paragraphs = doc.paragraphs[start_index+1:end_index]

#     # Create a new Word document to hold the entry
#     new_doc = docx.Document()

#     # Copy over the entry paragraphs
#     for para in entry_paragraphs:
#         # Add the paragraph to the new document
#         new_para = new_doc.add_paragraph()

#         # Copy over the paragraph text and formatting
#         for run in para.runs:
#             new_run = new_para.add_run(run.text)
#             new_run.bold = run.bold
#             new_run.italic = run.italic
#             new_run.underline = run.underline
#             new_run.font.color.rgb = run.font.color.rgb

#             # Copy over tracked changes
#             if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
#                 new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

#             # Copy over comments
#             for comment in doc.comments:
#                 if comment.reference._element == run._element:
#                     new_comment = new_para._element.add_comment(comment.author, comment.initial)
#                     new_comment.add_p(run.text)

#     # Save the new document
#     new_doc.save('entry.docx')



#start_text_index = find_start_text_index(word_document,18)

# placename = get_placename(word_document, start_text_index)

# end_text_index = find_end_text_index(word_document,start_text_index)

# placename_id = get_placename_id(word_document,end_text_index)

# filename = placename  +"." + placename_id + "." + "4"

# entry = get_entry_text(word_document, start_text_index, end_text_index)

# print('placename: ', placename)
# print('placename_id: ', placename_id)
# print(filename)
#print('start_text index: ', start_text_index)
# print('end_text_index: ', end_text_index)
# print(entry)
